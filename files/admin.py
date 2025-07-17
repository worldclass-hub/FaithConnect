from django.contrib import admin, messages
from django.utils.html import format_html
from django.template.response import TemplateResponse
from django.urls import reverse, path
from django.http import JsonResponse, HttpResponse
from django.shortcuts import redirect
from django.db.models import Q
from django.templatetags.static import static
from django.core.mail import send_mail
from io import BytesIO

# PDF
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from reportlab.lib import colors

# Word
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from PIL import Image
from docx.oxml import parse_xml  # ✅ Correct import
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


import os
import csv
import io
import csv
import datetime
from calendar import month_name
from collections import defaultdict
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors

from ckeditor.widgets import CKEditorWidget

from .models import (
    FileUpload, PDFUpload, Gallery, Hymn, Hymn_Content,
    FrenchHymn, YorubaHymn, IgboHymn, HausaHymn, ChineseHymn, GermanHymn,
    NewsletterSubscriber, DailyNewsletter,
    NewUpdate, UserProfile,
    ComingSoonPage, AboutPage, ContactMessage, AutoReplyMessage, PrayerRequest
)






















from django.contrib import admin
from django.utils.html import format_html
from django.templatetags.static import static
import os
from .models import FileUpload, PDFUpload, Gallery


@admin.register(FileUpload)
class FileUploadAdmin(admin.ModelAdmin):
    list_display = ('date', 'time', 'company_location', 'file_icon', 'file_url')

    def file_icon(self, obj):
        """Show an icon based on file extension (excluding PDF)"""
        if not obj.file_url:
            return "-"
        ext = os.path.splitext(obj.file_url)[1].lower()

        if ext in ['.jpg', '.jpeg', '.png', '.gif']:
            icon = 'image-icon.png'
        elif ext == '.txt':
            icon = 'text-icon.png'
        elif ext == '.zip':
            icon = 'zip-icon.png'
        else:
            icon = 'default-file-icon.png'

        icon_url = static(f'login-form/images/{icon}')
        return format_html('<img src="{}" width="30" style="object-fit: contain;" />', icon_url)

    file_icon.short_description = 'File Icon'


@admin.register(PDFUpload)
class PDFUploadAdmin(admin.ModelAdmin):
    list_display = ('company_location', 'info_id', 'pdf_url', 'date', 'time')


class GalleryAdmin(admin.ModelAdmin):
    list_display = ('title', 'image_display')

    def image_display(self, obj):
        return format_html('<img src="{}" style="width: 100px; height: auto;" />', obj.image_url)

    image_display.short_description = 'Image'


admin.site.register(Gallery, GalleryAdmin)















class HymnAdmin(admin.ModelAdmin):
    list_display = ('title', 'hymn_type',  'description')
    search_fields = ('title', 'description', 'hymn_type')

admin.site.register(Hymn, HymnAdmin)












class ChineseHymnAdmin(admin.ModelAdmin):
    list_display = ('title', 'hymn_type',  'description')
    search_fields = ('title', 'description', 'hymn_type')

admin.site.register(ChineseHymn, ChineseHymnAdmin)









class GermanHymnAdmin(admin.ModelAdmin):
    list_display = ('title', 'hymn_type',  'description')
    search_fields = ('title', 'description', 'hymn_type')

admin.site.register(GermanHymn, GermanHymnAdmin)












class FrenchHymnAdmin(admin.ModelAdmin):
    list_display = ('title', 'hymn_type',  'description')
    search_fields = ('title', 'description', 'hymn_type')

admin.site.register(FrenchHymn, FrenchHymnAdmin)

















class YorubaHymnAdmin(admin.ModelAdmin):
    list_display = ('title', 'hymn_type',  'description')
    search_fields = ('title', 'description', 'hymn_type')

admin.site.register(YorubaHymn, YorubaHymnAdmin)

















class IgboHymnAdmin(admin.ModelAdmin):
    list_display = ('title', 'hymn_type',  'description')
    search_fields = ('title', 'description', 'hymn_type')

admin.site.register(IgboHymn, IgboHymnAdmin)

















class HausaHymnAdmin(admin.ModelAdmin):
    list_display = ('title', 'hymn_type',  'description')
    search_fields = ('title', 'description', 'hymn_type')

admin.site.register(HausaHymn, HausaHymnAdmin)
















@admin.register(Hymn_Content)
class EnglishHymnAdmin(admin.ModelAdmin):
    list_display = ('created_at', 'short_lyrics')
    search_fields = ('lyrics',)

    def short_lyrics(self, obj):
        return (obj.lyrics[:75] + '...') if len(obj.lyrics) > 75 else obj.lyrics
    short_lyrics.short_description = 'Lyrics'


















@admin.register(ContactMessage)
class ContactMessageAdmin(admin.ModelAdmin):
    list_display = ('name', 'email', 'phone', 'subject', 'submitted_at')
    list_filter = ('submitted_at',)
    search_fields = ('name', 'email', 'subject', 'message')
    actions = ['mark_as_resolved']

    def mark_as_resolved(self, request, queryset):
        count = queryset.count()
        # This is just a placeholder action – it doesn't change anything yet
        messages.success(request, f"{count} message(s) marked as resolved.")
    
    mark_as_resolved.short_description = "Mark selected messages as resolved"
















@admin.register(AutoReplyMessage)
class AutoReplyMessageAdmin(admin.ModelAdmin):
    list_display = ('subject', 'message')
    search_fields = ('subject', 'message')
















from django.contrib import admin
from .models import AboutPage
from .forms import AboutPageForm

class AboutPageAdmin(admin.ModelAdmin):
    form = AboutPageForm

    fieldsets = (
        (None, {
            'fields': ('title', 'history', 'more_history', 'the_man', 'mission_statement', 'vision', 'value', 'mandate')
        }),
        ('Contact Information', {
            'fields': ('contact_address', 'phone', 'email')
        }),
        ('Image Upload', {
            'fields': ('image1_upload',)
        }),
    )

admin.site.register(AboutPage, AboutPageAdmin)

















# admin.py
from django.contrib import admin
from django.utils.html import format_html
from .models import ComingSoonPage
from .forms import ComingSoonForm

@admin.register(ComingSoonPage)
class ComingSoonPageAdmin(admin.ModelAdmin):
    form = ComingSoonForm
    list_display = ('id', 'note', 'updated_at', 'background_preview')
    readonly_fields = ('background_preview',)

    fieldsets = (
        (None, {
            'fields': ('note', 'background_upload',)
        }),
    )

    def background_preview(self, obj):
        if obj.background_image_url:
            return format_html('<img src="{}" style="max-height: 200px;">', obj.background_image_url)
        return "No image uploaded"

    background_preview.short_description = "Background Image Preview"
















@admin.register(UserProfile)
class UserProfileAdmin(admin.ModelAdmin):
    list_display = ('fullname', 'dob', 'gender', 'phone', 'gmail', 'profile_image_tag')
    list_filter = ('dob',)
    search_fields = ('fullname', 'gmail', 'phone')
    change_list_template = 'files/custom_content_wrapper.html'

    def profile_image_tag(self, obj):
        if obj.profile_image:
            return format_html(
                '<img src="{}" width="60" height="60" style="object-fit: cover; border-radius: 6px;" />',
                obj.profile_image.url
            )
        return "-"
    profile_image_tag.short_description = 'Profile Image'

    def changelist_view(self, request, extra_context=None):
        grouped_profiles = defaultdict(list)
        for profile in UserProfile.objects.all():
            month = profile.dob.strftime('%B') if profile.dob else 'Unknown'
            grouped_profiles[month].append(profile)

        months_order = list(month_name)[1:]
        sorted_profiles = {month: grouped_profiles.get(month, []) for month in months_order}

        extra_context = extra_context or {}
        extra_context['months'] = sorted_profiles
        extra_context['title'] = "User Profiles by Month"
        extra_context['app_label'] = 'files'
        extra_context['delete_url'] = self.get_delete_url()

        return super().changelist_view(request, extra_context=extra_context)

    def get_urls(self):
        urls = super().get_urls()
        custom_urls = [
            path('delete/<int:profile_id>/', self.admin_site.admin_view(self.delete_profile), name='delete_profile'),
            path('autocomplete/', self.admin_site.admin_view(self.autocomplete_search), name='userprofile_autocomplete'),
            path('export/csv/', self.admin_site.admin_view(self.export_profiles_csv), name='export_profiles_csv'),
            path('export/pdf/', self.admin_site.admin_view(self.export_profiles_pdf), name='export_profiles_pdf'),
        ]
        return custom_urls + urls

    def get_delete_url(self):
        return '/admin/files/userprofile/delete/'

    def delete_profile(self, request, profile_id):
        try:
            profile = UserProfile.objects.get(pk=profile_id)
            profile.delete()
            if request.headers.get('x-requested-with') == 'XMLHttpRequest':
                return JsonResponse({'status': 'success', 'message': 'Profile deleted successfully!'})
            else:
                self.message_user(request, "Profile deleted successfully!")
                return redirect('/admin/files/userprofile/')
        except UserProfile.DoesNotExist:
            if request.headers.get('x-requested-with') == 'XMLHttpRequest':
                return JsonResponse({'status': 'error', 'message': 'Profile not found!'}, status=404)
            else:
                self.message_user(request, "Profile not found!", level="error")
                return redirect('/admin/files/userprofile/')

    def autocomplete_search(self, request):
        query = request.GET.get('q', '')
        results = []

        if query:
            profiles = UserProfile.objects.filter(
                Q(fullname__icontains=query) |
                Q(gmail__icontains=query) |
                Q(phone__icontains=query)
            )
            results = [
                {
                    'id': profile.id,
                    'label': profile.fullname,
                    'month': profile.dob.strftime('%B') if profile.dob else 'Unknown',
                }
                for profile in profiles
            ]
        return JsonResponse(results, safe=False)

    def export_profiles_csv(self, request):
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="All Users BirthDay Profile.csv"'

        writer = csv.writer(response)
        writer.writerow(['S/N', 'Full Name', 'DOB', 'Gender', 'Phone', 'Gmail', 'Month'])

        for index, profile in enumerate(UserProfile.objects.all().order_by('dob'), start=1):
            month = profile.dob.strftime('%B') if profile.dob else 'Unknown'
            writer.writerow([
                index,
                profile.fullname,
                profile.dob,
                profile.gender,
                profile.phone,
                profile.gmail,
                month
            ])
        return response

    def export_profiles_pdf(self, request):
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="All Users BirthDay Profile.pdf"'

        buffer = io.BytesIO()
        p = canvas.Canvas(buffer, pagesize=A4)
        
        # Set PDF document title
        p.setTitle("All User Profiles with DOB")

        width, height = A4

        y = height - 40
        p.setFont("Helvetica-Bold", 14)
        p.drawString(100, y, "All User Profile With D.O.B")
        y -= 30

        p.setFillColor(colors.green)
        p.rect(40, y, width - 80, 20, fill=True)
        p.setFillColor(colors.white)
        p.setFont("Helvetica-Bold", 10)
        p.drawString(45, y + 5, "S/N")
        p.drawString(110, y + 5, "Full Name")
        p.drawString(230, y + 5, "DOB")
        p.drawString(310, y + 5, "Gender")
        p.drawString(400, y + 5, "Phone")
        p.drawString(500, y + 5, "Month")
        y -= 20

        p.setFont("Helvetica", 9)
        row_color = colors.white
        for index, profile in enumerate(UserProfile.objects.all().order_by('dob'), start=1):
            month = profile.dob.strftime('%B') if profile.dob else 'Unknown'

            # Alternating row colors
            if row_color == colors.white:
                row_color = colors.lightgrey
            else:
                row_color = colors.white

            p.setFillColor(row_color)
            p.rect(40, y, width - 80, 20, fill=True)

            p.setFillColor(colors.black)
            p.drawString(45, y + 5, str(index))
            p.drawString(110, y + 5, profile.fullname)
            p.drawString(230, y + 5, str(profile.dob))
            p.drawString(310, y + 5, profile.gender)
            p.drawString(400, y + 5, profile.phone)
            p.drawString(500, y + 5, month)

            y -= 25
            if y < 50:
                p.showPage()
                y = height - 40
                p.setFillColor(colors.green)
                p.rect(40, y, width - 80, 20, fill=True)
                p.setFillColor(colors.white)
                p.setFont("Helvetica-Bold", 10)
                p.drawString(45, y + 5, "S/N")
                p.drawString(110, y + 5, "Full Name")
                p.drawString(230, y + 5, "DOB")
                p.drawString(310, y + 5, "Gender")
                p.drawString(400, y + 5, "Phone")
                p.drawString(500, y + 5, "Month")
                y -= 20

        p.save()
        pdf = buffer.getvalue()
        buffer.close()
        response.write(pdf)
        return response




from django.contrib import admin
from django.utils.html import format_html
from .models import NewUpdate
from .forms import NewUpdateForm

@admin.register(NewUpdate)
class NewUpdateAdmin(admin.ModelAdmin):
    form = NewUpdateForm

    list_display = ('title', 'upload_date', 'user', 'image_tag')
    search_fields = ('title',)
    readonly_fields = ('image_url', 'image_tag')

    fieldsets = (
        (None, {
            'fields': ('title', 'content', 'user', 'image_upload', 'image_url', 'image_tag')
        }),
    )

    def image_tag(self, obj):
        if obj.image_url:
            return format_html(f'<img src="{obj.image_url}" style="height: 100px; border-radius: 8px; box-shadow: 0 2px 6px rgba(0,0,0,0.3);">')
        return "-"
    image_tag.short_description = 'Preview'














# Custom action to send a newsletter to all subscribers
@admin.action(description="Send selected newsletter to all subscribers")
def send_newsletter_to_all(modeladmin, request, queryset):
    for newsletter in queryset:
        subscribers = NewsletterSubscriber.objects.all()
        for subscriber in subscribers:
            send_mail(
                newsletter.subject,
                newsletter.body,
                'doxcela@gmail.com',  # Sender's email (make sure it matches your settings)
                [subscriber.email],
                fail_silently=False,
            )
    modeladmin.message_user(request, "Newsletter sent successfully.")

# Admin configuration for DailyNewsletter
class DailyNewsletterAdmin(admin.ModelAdmin):
    list_display = ('subject', 'date_posted')
    actions = [send_newsletter_to_all]  # Add custom action here

# Registering models in the admin panel
admin.site.register(NewsletterSubscriber)  # To show the list of subscribers
admin.site.register(DailyNewsletter, DailyNewsletterAdmin)  # To show the newsletter list and actions









from django.contrib import admin
from django.http import HttpResponse
from .models import PrayerRequest
from io import BytesIO
import csv, os, datetime
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.utils import ImageReader
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import parse_xml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

@admin.register(PrayerRequest)
class PrayerRequestAdmin(admin.ModelAdmin):
    list_display = ('name', 'email', 'phone', 'created_at')
    list_filter = ('created_at',)
    search_fields = ('name', 'email', 'phone', 'message')
    readonly_fields = ('name', 'email', 'phone', 'message', 'created_at')
    ordering = ('-created_at',)
    actions = ['export_as_csv', 'export_as_pdf', 'export_as_word']

    def has_add_permission(self, request):
        return False

    # ✅ CSV Export
    def export_as_csv(self, request, queryset):
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="prayer_requests.csv"'
        writer = csv.writer(response)
        writer.writerow(['Name', 'Email', 'Phone', 'Message', 'Created At'])
        for prayer in queryset:
            writer.writerow([prayer.name, prayer.email, prayer.phone, prayer.message, prayer.created_at])
        return response
    export_as_csv.short_description = "📤 Export selected to CSV"

    # ✅ PDF Export
    def export_as_pdf(self, request, queryset):
        buffer = BytesIO()
        p = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        y = height - 100
        page_number = 1

        logo_path = os.path.join('static', 'login-form', 'images', 'GMMI_LOGO.png')
        watermark_text = "GMMIConnect"

        def draw_header():
            nonlocal y
            try:
                logo = ImageReader(logo_path)
                p.drawImage(logo, 40, height - 80, width=60, height=60, mask='auto')
            except:
                pass
            p.setFont("Helvetica-Bold", 16)
            p.drawString(110, height - 60, "Prayer Requests Report")
            p.setFont("Helvetica", 10)
            p.drawRightString(width - 40, height - 60, datetime.datetime.now().strftime("%Y-%m-%d %H:%M"))
            y = height - 100

        def draw_footer():
            p.setFont("Helvetica-Oblique", 8)
            p.setFillColor(colors.grey)
            p.drawCentredString(width / 2.0, 20, f"Page {page_number}")
            p.saveState()
            p.setFont("Helvetica-Bold", 40)
            p.setFillColorRGB(0.9, 0.9, 0.9)
            p.translate(width / 2, height / 2)
            p.rotate(45)
            p.drawCentredString(0, 0, watermark_text)
            p.restoreState()

        draw_header()
        p.setFont("Helvetica", 11)

        for obj in queryset:
            if y < 120:
                draw_footer()
                p.showPage()
                page_number += 1
                draw_header()
                p.setFont("Helvetica", 11)

            p.setFillColor(colors.black)
            p.drawString(50, y, f"Name: {obj.name}")
            y -= 18
            p.drawString(50, y, f"Email: {obj.email}")
            y -= 18
            p.drawString(50, y, f"Phone: {obj.phone or 'N/A'}")
            y -= 18
            p.drawString(50, y, f"Date: {obj.created_at.strftime('%Y-%m-%d %H:%M')}")
            y -= 18
            p.drawString(50, y, "Message:")
            y -= 16

            for line in obj.message.splitlines():
                if y < 80:
                    draw_footer()
                    p.showPage()
                    page_number += 1
                    draw_header()
                    p.setFont("Helvetica", 11)
                p.drawString(70, y, line.strip())
                y -= 14

            y -= 20

        draw_footer()
        p.save()
        buffer.seek(0)
        return HttpResponse(buffer, content_type='application/pdf', headers={
            'Content-Disposition': 'attachment; filename="prayer_requests.pdf"'
        })
    export_as_pdf.short_description = "🧾 Export selected to PDF"

    # ✅ Word Export
    def export_as_word(self, request, queryset):
        document = Document()
        section = document.sections[0]
        header = section.header
        header_paragraph = header.paragraphs[0]

        logo_path = os.path.join('static', 'login-form', 'images', 'GMMI_LOGO.png')
        if os.path.exists(logo_path):
            run = header_paragraph.add_run()
            run.add_picture(logo_path, width=Inches(1.5))
            header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        watermark_paragraph = header.add_paragraph()
        watermark_run = watermark_paragraph.add_run()
        shape_xml = r"""
        <w:pict xmlns:v="urn:schemas-microsoft-com:vml"
                xmlns:w="urn:schemas-microsoft-com:office:word"
                xmlns:o="urn:schemas-microsoft-com:office:office">
            <v:shape id="WordArt1" o:spid="_x0000_s1025" type="#_x0000_t136"
                     style="position:absolute; margin-left:0; margin-top:0;
                            width:500pt; height:100pt; z-index:-251654144;
                            mso-wrap-edited:f; rotation:315"
                     fillcolor="#e6e6e6" stroked="f">
                <v:textpath style="font-family:Calibri; font-size:36pt"
                            on="t" string="GMMIConnect"/>
                <v:fill opacity="0.1"/>
            </v:shape>
        </w:pict>
        """
        watermark_element = parse_xml(shape_xml)
        watermark_run._r.append(watermark_element)

        document.add_paragraph()
        title = document.add_heading("Prayer Requests Report", level=1)
        title.runs[0].font.size = Pt(24)
        title.runs[0].font.name = "Calibri"

        date_paragraph = document.add_paragraph()
        date_run = date_paragraph.add_run(f"Date Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
        date_run.bold = True
        date_run.font.size = Pt(10)
        date_run.font.name = "Calibri"
        document.add_paragraph()

        for obj in queryset:
            header = document.add_paragraph()
            header_run = header.add_run("🧎🏽 Prayer Request")
            header_run.bold = True
            header_run.font.size = Pt(13)
            header_run.font.name = "Calibri"
            header_run.font.color.rgb = RGBColor(46, 116, 181)

            p1 = document.add_paragraph()
            r1 = p1.add_run(f"Name: {obj.name}")
            r1.bold = True
            r1.font.size = Pt(11)

            p2 = document.add_paragraph()
            r2 = p2.add_run(f"Email: {obj.email}")
            r2.font.size = Pt(11)

            p3 = document.add_paragraph()
            r3 = p3.add_run(f"Phone: {obj.phone or 'N/A'}")
            r3.font.size = Pt(11)

            p4 = document.add_paragraph()
            r4 = p4.add_run(f"Date: {obj.created_at.strftime('%Y-%m-%d %H:%M')}")
            r4.font.size = Pt(11)

            document.add_paragraph("Message:", style="Intense Quote")
            for line in obj.message.splitlines():
                msg_paragraph = document.add_paragraph()
                msg_run = msg_paragraph.add_run(line.strip())
                msg_run.italic = True
                msg_run.font.size = Pt(10)
                msg_run.font.color.rgb = RGBColor(100, 100, 100)

            document.add_paragraph("\n")

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)

        return HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': 'attachment; filename="prayer_requests.docx"'}
        )
    export_as_word.short_description = "📝 Export selected to Word (.docx)"







from django.contrib import admin
from .models import WomenMinistryLeader

@admin.register(WomenMinistryLeader)
class WomenMinistryLeaderAdmin(admin.ModelAdmin):
    list_display = ['name', 'title']






from .models import MenMinistryLeader

@admin.register(MenMinistryLeader)
class MenMinistryLeaderAdmin(admin.ModelAdmin):
    list_display = ['name', 'title']






from .models import YouthMinistryLeader

@admin.register(YouthMinistryLeader)
class YouthMinistryLeaderAdmin(admin.ModelAdmin):
    list_display = ['name', 'title']





from .models import EvangelismMinistryLeader

@admin.register(EvangelismMinistryLeader)
class EvangelismMinistryLeaderAdmin(admin.ModelAdmin):
    list_display = ['name', 'title']





from django.contrib import admin
from django.http import HttpResponse
from .models import ChildrenMinistryRegistration
from io import BytesIO
import csv, os, datetime
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.utils import ImageReader
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import parse_xml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

@admin.register(ChildrenMinistryRegistration)
class ChildrenMinistryRegistrationAdmin(admin.ModelAdmin):
    list_display = ('parent_name', 'phone', 'child_name', 'age_group', 'submitted_at')
    list_filter = ('submitted_at', 'age_group')
    search_fields = ('parent_name', 'child_name', 'phone')
    readonly_fields = ('parent_name', 'phone', 'child_name', 'age_group', 'message', 'submitted_at')
    ordering = ('-submitted_at',)
    actions = ['export_as_csv', 'export_as_pdf', 'export_as_word']

    def has_add_permission(self, request):
        return False

    def export_as_csv(self, request, queryset):
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="children_ministry_registrations.csv"'
        writer = csv.writer(response)
        writer.writerow(['Parent Name', 'Phone', 'Child Name', 'Age Group', 'Message', 'Submitted At'])
        for entry in queryset:
            writer.writerow([
                entry.parent_name, entry.phone, entry.child_name,
                entry.age_group, entry.message, entry.submitted_at
            ])
        return response
    export_as_csv.short_description = "📤 Export selected to CSV"

    def export_as_pdf(self, request, queryset):
        buffer = BytesIO()
        p = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        y = height - 100
        page_number = 1

        logo_path = os.path.join('static', 'login-form', 'images', 'GMMI_LOGO.png')
        watermark_text = "GMMIConnect"

        def draw_header():
            nonlocal y
            try:
                logo = ImageReader(logo_path)
                p.drawImage(logo, 40, height - 80, width=60, height=60, mask='auto')
            except:
                pass
            p.setFont("Helvetica-Bold", 16)
            p.drawString(110, height - 60, "Children Ministry Registrations")
            p.setFont("Helvetica", 10)
            p.drawRightString(width - 40, height - 60, datetime.datetime.now().strftime("%Y-%m-%d %H:%M"))
            y = height - 100

        def draw_footer():
            p.setFont("Helvetica-Oblique", 8)
            p.setFillColor(colors.grey)
            p.drawCentredString(width / 2.0, 20, f"Page {page_number}")
            p.saveState()
            p.setFont("Helvetica-Bold", 40)
            p.setFillColorRGB(0.9, 0.9, 0.9)
            p.translate(width / 2, height / 2)
            p.rotate(45)
            p.drawCentredString(0, 0, watermark_text)
            p.restoreState()

        draw_header()
        p.setFont("Helvetica", 11)

        for obj in queryset:
            if y < 120:
                draw_footer()
                p.showPage()
                page_number += 1
                draw_header()
                p.setFont("Helvetica", 11)

            p.setFillColor(colors.black)
            p.drawString(50, y, f"Parent: {obj.parent_name}")
            y -= 18
            p.drawString(50, y, f"Phone: {obj.phone}")
            y -= 18
            p.drawString(50, y, f"Child: {obj.child_name}")
            y -= 18
            p.drawString(50, y, f"Age Group: {obj.age_group}")
            y -= 18
            p.drawString(50, y, f"Date: {obj.submitted_at.strftime('%Y-%m-%d %H:%M')}")
            y -= 18
            p.drawString(50, y, "Message:")
            y -= 16

            for line in obj.message.splitlines():
                if y < 80:
                    draw_footer()
                    p.showPage()
                    page_number += 1
                    draw_header()
                    p.setFont("Helvetica", 11)
                p.drawString(70, y, line.strip())
                y -= 14

            y -= 20

        draw_footer()
        p.save()
        buffer.seek(0)
        return HttpResponse(buffer, content_type='application/pdf', headers={
            'Content-Disposition': 'attachment; filename="children_ministry_registrations.pdf"'
        })
    export_as_pdf.short_description = "🧾 Export selected to PDF"

    def export_as_word(self, request, queryset):
        document = Document()
        section = document.sections[0]
        header = section.header
        header_paragraph = header.paragraphs[0]

        logo_path = os.path.join('static', 'login-form', 'images', 'GMMI_LOGO.png')
        if os.path.exists(logo_path):
            run = header_paragraph.add_run()
            run.add_picture(logo_path, width=Inches(1.5))
            header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        watermark_paragraph = header.add_paragraph()
        watermark_run = watermark_paragraph.add_run()
        shape_xml = r"""
        <w:pict xmlns:v="urn:schemas-microsoft-com:vml"
                xmlns:w="urn:schemas-microsoft-com:office:word"
                xmlns:o="urn:schemas-microsoft-com:office:office">
            <v:shape id="WordArt1" o:spid="_x0000_s1025" type="#_x0000_t136"
                     style="position:absolute; margin-left:0; margin-top:0;
                            width:500pt; height:100pt; z-index:-251654144;
                            mso-wrap-edited:f; rotation:315"
                     fillcolor="#e6e6e6" stroked="f">
                <v:textpath style="font-family:Calibri; font-size:36pt"
                            on="t" string="GMMIConnect"/>
                <v:fill opacity="0.1"/>
            </v:shape>
        </w:pict>
        """
        watermark_element = parse_xml(shape_xml)
        watermark_run._r.append(watermark_element)

        document.add_paragraph()
        title = document.add_heading("Children Ministry Registrations", level=1)
        title.runs[0].font.size = Pt(24)
        title.runs[0].font.name = "Calibri"

        date_paragraph = document.add_paragraph()
        date_run = date_paragraph.add_run(f"Date Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
        date_run.bold = True
        date_run.font.size = Pt(10)
        date_run.font.name = "Calibri"
        document.add_paragraph()

        for obj in queryset:
            header = document.add_paragraph()
            header_run = header.add_run("🧒 Registration Entry")
            header_run.bold = True
            header_run.font.size = Pt(13)
            header_run.font.name = "Calibri"
            header_run.font.color.rgb = RGBColor(46, 116, 181)

            p1 = document.add_paragraph(f"Parent: {obj.parent_name}")
            p2 = document.add_paragraph(f"Phone: {obj.phone}")
            p3 = document.add_paragraph(f"Child: {obj.child_name}")
            p4 = document.add_paragraph(f"Age Group: {obj.age_group}")
            p5 = document.add_paragraph(f"Date: {obj.submitted_at.strftime('%Y-%m-%d %H:%M')}")

            document.add_paragraph("Message:", style="Intense Quote")
            for line in obj.message.splitlines():
                document.add_paragraph(line.strip(), style='BodyText')

            document.add_paragraph("\n")

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)

        return HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': 'attachment; filename="children_ministry_registrations.docx"'}
        )
    export_as_word.short_description = "📝 Export selected to Word (.docx)"




from django.utils.html import format_html
from django.contrib import admin
from .models import GMSOMSlide
from .forms import GMSOMSlideForm

@admin.register(GMSOMSlide)
class GMSOMSlideAdmin(admin.ModelAdmin):
    form = GMSOMSlideForm
    list_display = ['caption', 'image_preview']
    readonly_fields = ['image_preview', 'image_url']

    def image_preview(self, obj):
        if obj.image_url:
            return format_html('<img src="{}" style="height: 100px;" />', obj.image_url)
        return "-"
    image_preview.short_description = 'Preview'





















# from .models import Testimony


# @admin.register(Testimony)
# class TestimonyAdmin(admin.ModelAdmin):
#     list_display = ("name", "quote", "created_at")
#     search_fields = ("name", "quote")




from django.contrib import admin
from django.http import HttpResponse
from .models import SchoolOfMinistryRegistration
from io import BytesIO
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Inches
import datetime

@admin.register(SchoolOfMinistryRegistration)
class SchoolOfMinistryRegistrationAdmin(admin.ModelAdmin):
    list_display = ('full_name', 'email', 'phone', 'country', 'submitted_at')
    readonly_fields = ('submitted_at',)
    ordering = ('-submitted_at',)
    actions = ['export_as_pdf', 'export_as_word']

    def has_add_permission(self, request):
        return False  # Optional: disable admin adding from backend

    def export_as_pdf(self, request, queryset):
        from django.template.loader import render_to_string

        buffer = BytesIO()
        result = BytesIO()

        for reg in queryset:
            context = {'registration': reg}
            html = render_to_string("files/pdf_template.html", context)
            pisa_status = pisa.pisaDocument(BytesIO(html.encode("UTF-8")), dest=result)

        response = HttpResponse(result.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="school_registrations.pdf"'
        return response
    export_as_pdf.short_description = "🖨️ Export selected to PDF"

    def export_as_word(self, request, queryset):
        document = Document()
        document.add_heading('GMMI School of Ministry Registrations', 0)

        for reg in queryset:
            document.add_paragraph(f"Full Name: {reg.full_name}")
            document.add_paragraph(f"Email: {reg.email}")
            document.add_paragraph(f"Phone: {reg.phone}")
            document.add_paragraph(f"Gender: {reg.gender}")
            document.add_paragraph(f"Date of Birth: {reg.dob}")
            document.add_paragraph(f"Marital Status: {reg.marital_status}")
            document.add_paragraph(f"Country: {reg.country}")
            document.add_paragraph(f"State/City: {reg.state_city}")
            document.add_paragraph(f"Occupation: {reg.occupation}")
            document.add_paragraph(f"Church: {reg.church}")
            document.add_paragraph(f"Born Again: {reg.born_again}")
            document.add_paragraph(f"Baptized in Holy Ghost: {reg.holy_ghost}")
            document.add_paragraph(f"Reason: {reg.reason}")
            document.add_paragraph("")

            if reg.photo:
                try:
                    document.add_paragraph("Passport Photograph:")
                    document.add_picture(reg.photo.path, width=Inches(1.5))
                except:
                    document.add_paragraph("Photo not available.")

            document.add_paragraph("\n" + "-"*30 + "\n")

        word_stream = BytesIO()
        document.save(word_stream)
        word_stream.seek(0)

        return HttpResponse(
            word_stream.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': 'attachment; filename="school_registrations.docx"'}
        )
    export_as_word.short_description = "📄 Export selected to Word"


