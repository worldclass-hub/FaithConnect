from django.shortcuts import render, redirect
from django.contrib.auth import authenticate, login, logout
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.urls import reverse
from django.core.paginator import Paginator
from django.http import JsonResponse, HttpResponse
from django.db.models import Q
from django.views.decorators.csrf import csrf_exempt
from django.templatetags.static import static
from django.core.mail import send_mail, EmailMultiAlternatives
from django.contrib.auth.models import User
from email.mime.image import MIMEImage
from urllib.parse import urlparse, parse_qs
from django.conf import settings

from datetime import date, datetime
import os
import re
import uuid
import requests
from django.core.files.uploadedfile import SimpleUploadedFile
from .upload_to_supabase import upload_file_to_supabase
from .supabase_client import supabase  # assuming it's inside 'files' app


# prayer/views.py

from .forms import FileUploadForm, PDFUploadForm, GalleryUploadForm, UserProfileForm, PrayerRequestForm
from .models import (
    FileUpload, PDFUpload, Gallery, Hymn, FrenchHymn, HausaHymn, IgboHymn, YorubaHymn,
    ChineseHymn, GermanHymn, Hymn_Content, AboutPage, ContactMessage, AutoReplyMessage, ComingSoonPage,
    UserProfile, NewUpdate, NewsletterSubscriber, DailyNewsletter, Donation
)
from .models import NewUpdate  # import both models

def welcome_view(request):
    updates = NewUpdate.objects.all().order_by('-upload_date')
    # testimonies = Testimony.objects.all()

    return render(request, 'files/welcome.html', {
        'updates': updates,
        # 'testimonies': testimonies,
    })





def user_login(request):
    if request.method == "POST":
        username = request.POST["username"]
        password = request.POST["password"]
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            return redirect("excel_page")  # Redirect to home page after login
        else:
            messages.error(request, "Invalid username or password")
    return render(request, "files/login.html")



from .models import ComingSoonPage

def coming_soon(request):
    page = ComingSoonPage.objects.last()  # use .last() to get most recent
    return render(request, 'files/coming_soon.html', {'page': page})



















def signup_view(request):
    if request.method == 'POST':
        first_name = request.POST.get('first_name')
        last_name = request.POST.get('last_name')
        email = request.POST.get('email')
        username = request.POST.get('username')
        password = request.POST.get('password')
        confirm_password = request.POST.get('confirm_password')

        if password != confirm_password:
            messages.error(request, "Passwords do not match.")
            return redirect('signup')

        if User.objects.filter(username=username).exists():
            messages.error(request, "Username already exists.")
            return redirect('signup')

        if User.objects.filter(email=email).exists():
            messages.error(request, "Email already exists.")
            return redirect('signup')

        # ✅ Create and login the user
        user = User.objects.create_user(
            username=username,
            password=password,
            email=email,
            first_name=first_name,
            last_name=last_name
        )
        user.save()

        login(request, user)  # ✅ Log the user in

        messages.success(request, "Account created and logged in successfully!")
        return redirect('excel_page')

    return render(request, 'files/signup.html')


















from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from .forms import FileUploadForm
from .upload_to_supabase import upload_file_to_supabase
from .models import FileUpload
import os

@login_required
def home(request):
    if request.method == "POST":
        form = FileUploadForm(request.POST, request.FILES)
        if form.is_valid():
            instance = form.save(commit=False)
            uploaded_file = request.FILES.get("uploaded_file")
            youtube_url = form.cleaned_data.get("youtube_url")
            temp_path = None
            supabase_url = None

            # ✅ Handle file upload
            if uploaded_file:
                print(f"🎯 File selected: {uploaded_file.name}")
                try:
                    # Save temporarily
                    temp_path = f"/tmp/{uploaded_file.name}"
                    with open(temp_path, "wb+") as temp_file:
                        for chunk in uploaded_file.chunks():
                            temp_file.write(chunk)

                    # Upload to Supabase
                    supabase_url = upload_file_to_supabase(temp_path, "user-uploads")
                    print("✅ Supabase URL:", supabase_url)

                    if supabase_url:
                        instance.file_url = supabase_url
                    else:
                        print("❌ Upload failed.")

                except Exception as e:
                    print("❌ Upload error:", e)

                finally:
                    # Always clean up
                    if temp_path and os.path.exists(temp_path):
                        os.remove(temp_path)
                    else:
                        print(f"⚠️ Temp file missing or already removed: {temp_path}")

            # ✅ Manual validation
            if not instance.file_url and not youtube_url:
                form.add_error(None, "You must upload a file or provide a YouTube URL.")
                print("❌ No file or YouTube URL provided")
                return render(request, "files/home.html", {"form": form})

            # ✅ Block PDF
            try:
                if instance.file_url and instance.file_extension() == '.pdf':
                    form.add_error(None, "PDF files are not allowed.")
                    print("❌ PDF is not allowed")
                    return render(request, "files/home.html", {"form": form})
            except Exception as e:
                print("⚠️ Could not check file extension:", e)

            # ✅ Save instance
            instance.youtube_url = youtube_url
            instance.save()
            print("✅ File saved successfully!")
            return redirect("excel_page")
        else:
            print("❌ FORM IS NOT VALID")
            print(form.errors)
    else:
        form = FileUploadForm()

    return render(request, "files/home.html", {"form": form})
















from django.views.decorators.csrf import csrf_exempt
from django.shortcuts import render
from django.http import JsonResponse
from .models import FileUpload, NewsletterSubscriber, UserProfile, NewUpdate
from .forms import UserProfileForm
from datetime import datetime
from .upload_to_supabase import upload_file_to_supabase
import os


def excel_page(request):
    allowed_image_exts = ['.jpg', '.jpeg', '.png', '.gif']
    allowed_video_exts = ['.mp4', '.webm', '.ogg']
    allowed_audio_exts = ['.mp3', '.wav', '.aac']
    allowed_doc_exts = ['.doc', '.docx']
    allowed_ppt_exts = ['.ppt', '.pptx']
    allowed_excel_exts = ['.xls', '.xlsx']
    allowed_zip_exts = ['.zip', '.rar']

    if request.method == 'POST' and request.headers.get('x-requested-with') == 'XMLHttpRequest':

        # 🎯 Profile Form
        if 'fullname' in request.POST and request.FILES.get('profile_image'):
            if request.user.is_authenticated:
                if UserProfile.objects.filter(user=request.user).exists():
                    return JsonResponse({'status': 'error', 'message': "Profile already exists."})

                form = UserProfileForm(request.POST, request.FILES)
                if form.is_valid():
                    profile = form.save(commit=False)
                    profile.user = request.user
                    profile.save()
                    return JsonResponse({'status': 'success'})
                else:
                    return JsonResponse({'status': 'error', 'errors': form.errors}, status=400)
            else:
                return JsonResponse({'status': 'error', 'message': "You must be logged in."}, status=400)

        # 📬 Newsletter Subscription
        elif 'email' in request.POST and not request.FILES:
            email = request.POST.get('email')
            if NewsletterSubscriber.objects.filter(email=email).exists():
                return JsonResponse({'status': 'error', 'message': 'Email already subscribed.'})
            NewsletterSubscriber.objects.create(email=email)
            return JsonResponse({'status': 'success'})

        # 📁 File or 🎥 YouTube Upload
        uploaded_file = request.FILES.get('uploaded_file')
        youtube_url = request.POST.get('youtube_url', '').strip()
        date = request.POST.get('date')
        time = request.POST.get('time')
        company_location = request.POST.get('company_location')

        if not (date and time and company_location):
            return JsonResponse({'status': 'error', 'message': 'Date, time, and company location are required.'})
        if not uploaded_file and not youtube_url:
            return JsonResponse({'status': 'error', 'message': 'Please upload a file or provide a YouTube video URL.'})

        try:
            parsed_date = datetime.strptime(date, '%Y-%m-%d').date()
            parsed_time = datetime.strptime(time, '%H:%M').time()
        except ValueError:
            return JsonResponse({'status': 'error', 'message': 'Invalid date or time format.'})

        file_url = None
        if uploaded_file:
            try:
                temp_path = f"/tmp/{uploaded_file.name}"
                with open(temp_path, "wb+") as temp_file:
                    for chunk in uploaded_file.chunks():
                        temp_file.write(chunk)

                print("🎯 Uploading:", uploaded_file.name)
                file_url = upload_file_to_supabase(temp_path, "user-uploads")
                os.remove(temp_path)
                print("✅ Uploaded. Supabase URL:", file_url)

            except Exception as e:
                print("❌ Upload error:", e)
                return JsonResponse({'status': 'error', 'message': f"Upload failed: {str(e)}"})

        FileUpload.objects.create(
            file_url=file_url,
            youtube_url=youtube_url if youtube_url else None,
            date=parsed_date,
            time=parsed_time,
            company_location=company_location
        )

        return JsonResponse({'status': 'success', 'message': 'File or video uploaded successfully.'})

    # --- GET Request ---
    files = FileUpload.objects.all().order_by('-date', '-time')
    selected_date = request.GET.get('date')
    if selected_date:
        selected_date = datetime.strptime(selected_date, '%Y-%m-%d')
        files = files.filter(date=selected_date.date())

    user_has_profile = False
    if request.user.is_authenticated:
        user_has_profile = UserProfile.objects.filter(user=request.user).exists()

    updates = NewUpdate.objects.all().order_by('-upload_date')

    return render(request, 'files/excel_page.html', {
        'files': files,
        'user_has_profile': user_has_profile,
        'updates': updates,
        'allowed_image_exts': allowed_image_exts,
        'allowed_video_exts': allowed_video_exts,
        'allowed_audio_exts': allowed_audio_exts,
        'allowed_doc_exts': allowed_doc_exts,
        'allowed_ppt_exts': allowed_ppt_exts,
        'allowed_excel_exts': allowed_excel_exts,
        'allowed_zip_exts': allowed_zip_exts,
    })


# ✅ Optional (for when user closes modal and we want to suppress it again)
@csrf_exempt
def never_show_modal(request):
    if request.method == 'POST' and request.headers.get('x-requested-with') == 'XMLHttpRequest':
        email = request.POST.get('email')
        if email:
            NewsletterSubscriber.objects.filter(email=email).update(has_closed_modal=True)
            return JsonResponse({'status': 'success'})
    return JsonResponse({'status': 'error'})

















from .upload_to_supabase import upload_file_to_supabase
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from .forms import PDFUploadForm
from .models import PDFUpload
from .upload_to_supabase import upload_file_to_supabase
import os

@login_required
def upload_pdf(request):
    if request.method == 'POST':
        form = PDFUploadForm(request.POST, request.FILES)
        files = request.FILES.getlist("pdf_files")  # ✅ MULTIPLE files manually

        print("📥 REQUEST.FILES:", request.FILES)
        print("📦 getlist(pdf_files):", files)

        if not files:
            form.add_error(None, "No PDF files were submitted.")
            return render(request, 'files/upload_pdf.html', {'form': form})

        if form.is_valid():
            image = request.FILES.get('image', None)

            for file in files:
                if not file.name.lower().endswith(".pdf"):
                    print("❌ Not a PDF:", file.name)
                    continue

                temp_path = f"/tmp/{file.name}"
                with open(temp_path, "wb+") as temp_file:
                    for chunk in file.chunks():
                        temp_file.write(chunk)

                pdf_url = upload_file_to_supabase(temp_path, "user-uploads")

                image_url = None
                if image:
                    image_path = f"/tmp/{image.name}"
                    with open(image_path, "wb+") as img_file:
                        for chunk in image.chunks():
                            img_file.write(chunk)
                    image_url = upload_file_to_supabase(image_path, "user-uploads")
                    os.remove(image_path)

                PDFUpload.objects.create(
                    company_location=form.cleaned_data['company_location'],
                    info_id=form.cleaned_data.get('info_id', ''),
                    pdf_url=pdf_url,
                    date=form.cleaned_data['date'],
                    time=form.cleaned_data['time'],
                    image_url=image_url or 'https://cdn-icons-png.flaticon.com/512/337/337946.png'
                )

                os.remove(temp_path)

            return redirect('pdf_document')
        else:
            print("❌ Form Errors:", form.errors)
    else:
        form = PDFUploadForm()

    return render(request, 'files/upload_pdf.html', {'form': form})










from django.core.paginator import Paginator
from django.http import JsonResponse
from django.shortcuts import render
from .models import PDFUpload, UserProfile, NewUpdate
from .forms import UserProfileForm
from datetime import datetime


def pdf_document(request):
    # ✅ Handle AJAX profile form submission
    if request.method == 'POST' and request.headers.get('x-requested-with') == 'XMLHttpRequest':
        if 'fullname' in request.POST and request.FILES.get('profile_image'):
            if request.user.is_authenticated:
                if UserProfile.objects.filter(user=request.user).exists():
                    return JsonResponse({'status': 'error', 'message': "Profile already exists."}, status=400)

                form = UserProfileForm(request.POST, request.FILES)
                if form.is_valid():
                    profile = form.save(commit=False)
                    profile.user = request.user
                    profile.save()
                    return JsonResponse({'status': 'success'})
                else:
                    return JsonResponse({'status': 'error', 'errors': form.errors}, status=400)
            else:
                return JsonResponse({'status': 'error', 'message': "You must be logged in."}, status=400)

    # ✅ Fetch all PDFs
    pdf_list = PDFUpload.objects.all().order_by('-id')
    selected_date = request.GET.get('date', None)
    if selected_date:
        selected_date = datetime.strptime(selected_date, '%Y-%m-%d')
        pdf_list = pdf_list.filter(date=selected_date.date())

    paginator = Paginator(pdf_list, 6)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    recent_pdfs = PDFUpload.objects.all().order_by('-id')[:5]

    user_has_profile = False
    if request.user.is_authenticated:
        user_has_profile = UserProfile.objects.filter(user=request.user).exists()

    updates = NewUpdate.objects.all().order_by('-upload_date')

    return render(request, "files/pdf_document.html", {
        'page_obj': page_obj,
        'recent_pdfs': recent_pdfs,
        'user_has_profile': user_has_profile,
        'updates': updates,
    })



def user_logout(request):
    logout(request)
    return redirect("login")














from django.http import JsonResponse
from datetime import date
from urllib.parse import urlparse, parse_qs
from .models import FileUpload, PDFUpload


def extract_youtube_id(url):
    try:
        parsed_url = urlparse(url)
        netloc = parsed_url.netloc.lower()

        if 'youtube.com' in netloc:
            if parsed_url.path.startswith('/shorts/'):
                return parsed_url.path.split('/')[2]
            else:
                query = parse_qs(parsed_url.query)
                return query.get('v', [None])[0]
        elif 'youtu.be' in netloc:
            return parsed_url.path.lstrip('/')
    except Exception:
        return None
    return None


def get_file_for_date(request, year, month, day):
    selected_date = date(year, month, day)
    files = FileUpload.objects.filter(date=selected_date).order_by('-date', '-time')
    pdf_files = PDFUpload.objects.filter(date=selected_date).order_by('-date', '-time')
    all_files = list(files) + list(pdf_files)

    file_data = []

    for file in all_files:
        file_url = ''
        file_extension = ''
        youtube_url = None
        youtube_id = None

        if isinstance(file, FileUpload):
            if file.file_url:
                file_url = file.file_url
                file_extension = file.file_extension().lstrip(".") if file.file_extension() else ''
            elif file.youtube_url:
                youtube_url = file.youtube_url
                youtube_id = extract_youtube_id(youtube_url)
                file_extension = 'youtube'
        elif isinstance(file, PDFUpload):
            file_url = file.pdf_url or ''
            file_extension = file.file_extension().lstrip(".") if file.file_extension() else ''

        file_data.append({
            "file_url": file_url,  # ✅ For Excel calendar
            "uploaded_file_url": file_url,  # ✅ For Search calendar
            "file_extension": file_extension,
            "company_location": file.company_location,
            "date": file.date.strftime("%Y-%m-%d"),
            "time": file.time.strftime("%I:%M %p"),
            "image_url": getattr(file, 'image_url', None),
            "info_id": getattr(file, 'info_id', "No Info ID"),
            "youtube_url": youtube_url,
            "youtube_id": youtube_id,
        })

    return JsonResponse({"files": file_data})










from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from .models import Gallery
from .upload_to_supabase import upload_file_to_supabase
import os

@login_required
def upload_image(request):
    if request.method == 'POST':
        title = request.POST.get('title')
        files = request.FILES.getlist('uploaded_files')

        print("📥 POST title:", title)
        print("📦 Uploaded files:", files)

        if not title:
            print("❌ Missing title")
            return render(request, 'files/gallery_upload.html', {
                'error': 'Please enter a title.'
            })

        if not files:
            print("❌ No files submitted")
            return render(request, 'files/gallery_upload.html', {
                'error': 'Please select at least one image.'
            })

        for file in files:
            try:
                # Save temporarily
                temp_path = f"/tmp/{file.name}"
                with open(temp_path, "wb+") as temp_file:
                    for chunk in file.chunks():
                        temp_file.write(chunk)

                # Upload to Supabase
                image_url = upload_file_to_supabase(temp_path, "user-uploads")

                # Save to DB
                Gallery.objects.create(title=title, image_url=image_url)
                print("✅ Uploaded and saved:", file.name)

                # Delete temp file
                os.remove(temp_path)

            except Exception as e:
                print(f"🔥 Error uploading {file.name}: {e}")
                return render(request, 'files/gallery_upload.html', {
                    'error': f'Error uploading {file.name}. Try again.'
                })

        return redirect('lookbook')

    return render(request, 'files/gallery_upload.html')







@login_required
def lookbook(request):
    gallery_images = Gallery.objects.all().order_by('-id')

    paginator = Paginator(gallery_images, 7)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    updates = NewUpdate.objects.all().order_by('-upload_date')

    return render(request, 'files/lookbook.html', {
        'page_obj': page_obj,
        'updates': updates,
    })




















# views.py
def hymn_list(request):
    hymns = Hymn.objects.all()
    updates = NewUpdate.objects.all().order_by('-upload_date')
    return render(request, 'files/hymn_list.html', {'hymns': hymns, 'updates': updates})


def hymn_detail(request, hymn_id):
    hymn = Hymn.objects.get(id=hymn_id)
    updates = NewUpdate.objects.all().order_by('-upload_date')
    return render(request, 'files/hymn_detail.html', {'hymn': hymn, 'updates': updates})


def french_hymn_list(request):
    hymns = FrenchHymn.objects.all()
    updates = NewUpdate.objects.all().order_by('-upload_date')
    return render(request, 'files/french_hymn_list.html', {'hymns': hymns, 'updates': updates})


def french_hymn_detail(request, hymn_id):
    hymn = FrenchHymn.objects.get(id=hymn_id)
    updates = NewUpdate.objects.all().order_by('-upload_date')
    return render(request, 'files/french_hymn_detail.html', {'hymn': hymn, 'updates': updates})


def yoruba_hymn_list(request):
    hymns = YorubaHymn.objects.all()
    updates = NewUpdate.objects.all().order_by('-upload_date')
    return render(request, 'files/yoruba_hymn_list.html', {'hymns': hymns, 'updates': updates})


def yoruba_hymn_detail(request, hymn_id):
    hymn = YorubaHymn.objects.get(id=hymn_id)
    updates = NewUpdate.objects.all().order_by('-upload_date')
    return render(request, 'files/yoruba_hymn_detail.html', {'hymn': hymn, 'updates': updates})


def igbo_hymn_list(request):
    hymns = IgboHymn.objects.all()
    updates = NewUpdate.objects.all().order_by('-upload_date')
    return render(request, 'files/igbo_hymn_list.html', {'hymns': hymns, 'updates': updates})


def igbo_hymn_detail(request, hymn_id):
    hymn = IgboHymn.objects.get(id=hymn_id)
    updates = NewUpdate.objects.all().order_by('-upload_date')
    return render(request, 'files/igbo_hymn_detail.html', {'hymn': hymn, 'updates': updates})


def hausa_hymn_list(request):
    hymns = HausaHymn.objects.all()
    updates = NewUpdate.objects.all().order_by('-upload_date')
    return render(request, 'files/hausa_hymn_list.html', {'hymns': hymns, 'updates': updates})


def hausa_hymn_detail(request, hymn_id):
    hymn = HausaHymn.objects.get(id=hymn_id)
    updates = NewUpdate.objects.all().order_by('-upload_date')
    return render(request, 'files/hausa_hymn_detail.html', {'hymn': hymn, 'updates': updates})



def hymn_content(request):
    hymns = Hymn_Content.objects.all()
    updates = NewUpdate.objects.all().order_by('-upload_date')
    return render(request, 'files/hymn_content.html', {'hymns': hymns, 'updates': updates})



# Chinese Hymn Views
def chinese_hymn_list(request):
    hymns = ChineseHymn.objects.all()
    updates = NewUpdate.objects.all().order_by('-upload_date')
    return render(request, 'files/chinese_hymn_list.html', {'hymns': hymns, 'updates': updates})

def chinese_hymn_detail(request, hymn_id):
    hymn = ChineseHymn.objects.get(id=hymn_id)
    updates = NewUpdate.objects.all().order_by('-upload_date')
    return render(request, 'files/chinese_hymn_detail.html', {'hymn': hymn, 'updates': updates})


# German Hymn Views
def german_hymn_list(request):
    hymns = GermanHymn.objects.all()
    updates = NewUpdate.objects.all().order_by('-upload_date')
    return render(request, 'files/german_hymn_list.html', {'hymns': hymns, 'updates': updates})

def german_hymn_detail(request, hymn_id):
    hymn = GermanHymn.objects.get(id=hymn_id)
    updates = NewUpdate.objects.all().order_by('-upload_date')
    return render(request, 'files/german_hymn_detail.html', {'hymn': hymn, 'updates': updates})


















from django.http import JsonResponse
from django.shortcuts import render
from django.core.mail import EmailMultiAlternatives
from django.utils.timezone import now
from django.views.decorators.csrf import csrf_exempt
from .models import NewUpdate, NewsletterSubscriber, UserProfile
from datetime import datetime


def contact(request):
    if request.method == 'POST' and request.headers.get('x-requested-with') == 'XMLHttpRequest':
        email = request.POST.get('email')
        name = 'Subscriber'

        # 🔎 Try to get full name if logged in
        if request.user.is_authenticated:
            profile = UserProfile.objects.filter(user=request.user).first()
            if profile and profile.fullname:
                name = profile.fullname
            elif request.user.get_full_name():
                name = request.user.get_full_name()
            elif request.user.username:
                name = request.user.username

        if email:
            if NewsletterSubscriber.objects.filter(email=email).exists():
                return JsonResponse({
                    'status': 'error',
                    'message': 'Email already subscribed.',
                    'email': email
                })

            NewsletterSubscriber.objects.create(email=email)

            subject = "Thanks for subscribing to our newsletter!"
            text_body = (
                f"Hi {name},\n\n"
                "Thanks for subscribing to our newsletter! "
                "You’re now part of our community and will be the first to receive the latest updates, exclusive content, and special offers.\n\n"
                "We’re excited to have you with us!"
            )
            html_body = text_body.replace('\n', '<br>')

            html_email = f"""
            <div style="font-family: Arial, sans-serif; max-width: 600px; margin: auto; background: #ffffff; border-radius: 10px; padding: 30px; border: 1px solid #eee;">
                <div style="text-align: center; margin-bottom: 20px;">
                    <img src="https://i.imgur.com/yEFgd2V.png" alt="Admin Logo" style="height: 60px;">
                </div>
                <h2 style="color: #1d3557;">Hello {name},</h2>
                <p style="font-size: 16px; color: #333;">{html_body}</p>
                <div style="text-align: center; margin: 30px 0;">
                    <a href="https://gmmi.up.railway.app" style="background: #1d3557; color: white; text-decoration: none; padding: 12px 25px; border-radius: 30px; font-weight: bold;">
                        Visit Our Website
                    </a>
                </div>
                <hr style="border: none; border-top: 1px solid #eee;">
                <div style="text-align: center; margin-top: 20px;">
                    <p style="color: #888; font-size: 14px;">Stay connected with us</p>
                    <div>
                        <a href="https://facebook.com/yourpage" style="margin: 0 5px;">
                            <img src="https://cdn-icons-png.flaticon.com/512/733/733547.png" alt="Facebook" style="height: 24px;">
                        </a>
                        <a href="https://twitter.com/yourpage" style="margin: 0 5px;">
                            <img src="https://cdn-icons-png.flaticon.com/512/733/733579.png" alt="Twitter" style="height: 24px;">
                        </a>
                        <a href="https://instagram.com/yourpage" style="margin: 0 5px;">
                            <img src="https://cdn-icons-png.flaticon.com/512/2111/2111463.png" alt="Instagram" style="height: 24px;">
                        </a>
                        <a href="https://wa.me/2349057147497?text=Hi%2C%20I%20am%20contacting%20you%20from%20Doxcela" target="_blank" style="margin: 0 5px;">
                            <img src="https://cdn-icons-png.flaticon.com/512/733/733585.png" alt="WhatsApp" style="height: 24px;">
                        </a>
                    </div>
                    <p style="color: #aaa; font-size: 12px; margin-top: 10px;">&copy; {now().year} Doxcela. All rights reserved.</p>
                </div>
            </div>
            """

            # ✉️ Send the email
            reply_email = EmailMultiAlternatives(subject, text_body, 'globalmandateministryinc@gmail.com', [email])
            reply_email.attach_alternative(html_email, "text/html")
            reply_email.send()

            return JsonResponse({'status': 'success', 'show_modal': True, 'email': email})

        return JsonResponse({'status': 'error', 'message': 'Email required.'})

    # GET: load contact page with updates
    updates = NewUpdate.objects.all().order_by('-upload_date')
    return render(request, 'files/contact.html', {'updates': updates})

















def submit_contact(request):
    if request.method == 'POST':
        name = request.POST.get('name')
        email = request.POST.get('email')
        phone = request.POST.get('phone')
        subject = request.POST.get('subject')
        message = request.POST.get('message')
        image = request.FILES.get('image')

        # Save to database
        contact = ContactMessage.objects.create(
            name=name,
            email=email,
            phone=phone,
            subject=subject,
            message=message,
            image=image
        )

        # Format message
        formatted_message = message.replace('\n', '<br>')

        image_html = ""
        if image:
            image_html = '''
                <img src="cid:user_image" 
                     style="width: 60px; height: 60px; border-radius: 50%; object-fit: cover; margin-right: 10px; margin-top: 20px;">
            '''

        html_content = f"""
        <div style="font-family: Arial, sans-serif;">
            <h2 style="color: #1d3557;">New Contact Message</h2>
            <div style="display: flex; align-items: center;">
                {image_html}
                <div>
                    <p><strong>Name:</strong> {name}</p>
                    <p><strong>Email:</strong> {email}</p>
                    <p><strong>Phone:</strong> {phone}</p>
                </div>
            </div>
            <br>
            <p><strong>Subject:</strong> {subject}</p>
            <p><strong>Message:</strong><br>{formatted_message}</p>
        </div>
        """

        subject_line = f"New Contact Message: {subject}"
        from_email = settings.DEFAULT_FROM_EMAIL
        to_email = [settings.ADMIN_EMAIL]

        email_message = EmailMultiAlternatives(subject_line, '', from_email, to_email)
        email_message.attach_alternative(html_content, "text/html")

        # 🛠 Attach image to email
        if image:
            image.open()
            img = MIMEImage(image.read(), _subtype=image.content_type.split('/')[-1])
            img.add_header('Content-ID', '<user_image>')
            img.add_header('Content-Disposition', 'inline', filename=image.name)
            email_message.attach(img)

        email_message.send()

        # Auto-reply
        auto_reply = AutoReplyMessage.objects.first()
        if auto_reply:
            reply_subject = auto_reply.subject
            reply_body = auto_reply.message.format(name=name)
        else:
            reply_subject = "Thanks for contacting us"
            reply_body = f"Hi {name},\n\nThanks for reaching out. We'll get back to you shortly.\n\nCheers!"

        # Replace new lines with <br> in the auto-reply message
        reply_body_html = reply_body.replace('\n', '<br>')

        reply_html = f"""
        <div style="font-family: Arial, sans-serif; max-width: 600px; margin: auto; background: #ffffff; border-radius: 10px; padding: 30px; border: 1px solid #eee;">
            <div style="text-align: center; margin-bottom: 20px;">
                <img src="https://i.imgur.com/yEFgd2V.png" alt="Admin Logo" style="height: 60px;">
            </div>
            <h2 style="color: #1d3557;">Hello {name},</h2>
            <p style="font-size: 16px; color: #333;">
                {reply_body_html}
            </p>
            <div style="text-align: center; margin: 30px 0;">
                <a href="https://gmmi.up.railway.app" style="background: #1d3557; color: white; text-decoration: none; padding: 12px 25px; border-radius: 30px; font-weight: bold;">
                    Visit Our Website
                </a>
            </div>
            <hr style="border: none; border-top: 1px solid #eee;">
            <div style="text-align: center; margin-top: 20px;">
                <p style="color: #888; font-size: 14px;">Stay connected with us</p>
                <div>
                    <a href="https://facebook.com/yourpage" style="margin: 0 5px;">
                        <img src="https://cdn-icons-png.flaticon.com/512/733/733547.png" alt="Facebook" style="height: 24px;">
                    </a>
                    <a href="https://twitter.com/yourpage" style="margin: 0 5px;">
                        <img src="https://cdn-icons-png.flaticon.com/512/733/733579.png" alt="Twitter" style="height: 24px;">
                    </a>
                    <a href="https://instagram.com/yourpage" style="margin: 0 5px;">
                        <img src="https://cdn-icons-png.flaticon.com/512/2111/2111463.png" alt="Instagram" style="height: 24px;">
                    </a>
                   <a href="https://wa.me/2349057147497?text=Hi%2C%20I%20am%20contacting%20you%20from%20Doxcela" target="_blank" style="margin: 0 5px;">
                        <img src="https://cdn-icons-png.flaticon.com/512/733/733585.png" alt="WhatsApp" style="height: 24px;">
                    </a>


                </div>
                <p style="color: #aaa; font-size: 12px; margin-top: 10px;">&copy; {datetime.now().year} Doxcela. All rights reserved.</p>
            </div>
        </div>
        """

        # Send auto-reply email
        reply_email = EmailMultiAlternatives(reply_subject, reply_body, from_email, [email])
        reply_email.attach_alternative(reply_html, "text/html")
        reply_email.send()

        return JsonResponse({'status': 'success'})

    return JsonResponse({'status': 'fail'}, status=400)

















from django.shortcuts import render
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.core.mail import EmailMultiAlternatives
from datetime import datetime

from .models import AboutPage, NewsletterSubscriber, NewUpdate
from .upload_to_supabase import upload_file_to_supabase

def about_view(request):
    about = AboutPage.objects.first()

    # Handle email subscription (AJAX)
    if request.method == 'POST' and request.headers.get('x-requested-with') == 'XMLHttpRequest':
        email = request.POST.get('email')
        name = 'Subscriber'

        if request.user.is_authenticated:
            if request.user.get_full_name():
                name = request.user.get_full_name()
            elif request.user.username:
                name = request.user.username

        if email:
            if NewsletterSubscriber.objects.filter(email=email).exists():
                return JsonResponse({'status': 'error', 'message': 'Email already subscribed.', 'email': email})

            NewsletterSubscriber.objects.create(email=email)

            reply_subject = "Thanks for subscribing to our newsletter!"
            reply_body = (
                "Thanks for subscribing to our newsletter! "
                "You’re now part of our community and will be the first to receive the latest updates, exclusive content, and special offers.\n\n"
                "We’re excited to have you with us!"
            )
            reply_body_html = reply_body.replace('\n', '<br>')

            reply_html = f"""
            <div style="font-family: Arial, sans-serif; max-width: 600px; margin: auto; background: #ffffff; border-radius: 10px; padding: 30px; border: 1px solid #eee;">
                <div style="text-align: center; margin-bottom: 20px;">
                    <img src="https://i.imgur.com/yEFgd2V.png" alt="Admin Logo" style="height: 60px;">
                </div>
                <h2 style="color: #1d3557;">Hello {name},</h2>
                <p style="font-size: 16px; color: #333;">{reply_body_html}</p>
                <div style="text-align: center; margin: 30px 0;">
                    <a href="https://gmmi.up.railway.app" style="background: #1d3557; color: white; text-decoration: none; padding: 12px 25px; border-radius: 30px; font-weight: bold;">
                        Visit Our Website
                    </a>
                </div>
                <hr style="border: none; border-top: 1px solid #eee;">
                <div style="text-align: center; margin-top: 20px;">
                    <p style="color: #888; font-size: 14px;">Stay connected with us</p>
                    <div>
                        <a href="https://facebook.com/yourpage" style="margin: 0 5px;">
                            <img src="https://cdn-icons-png.flaticon.com/512/733/733547.png" alt="Facebook" style="height: 24px;">
                        </a>
                        <a href="https://twitter.com/yourpage" style="margin: 0 5px;">
                            <img src="https://cdn-icons-png.flaticon.com/512/733/733579.png" alt="Twitter" style="height: 24px;">
                        </a>
                        <a href="https://instagram.com/yourpage" style="margin: 0 5px;">
                            <img src="https://cdn-icons-png.flaticon.com/512/2111/2111463.png" alt="Instagram" style="height: 24px;">
                        </a>
                        <a href="https://wa.me/2349057147497?text=Hi%2C%20I%20am%20contacting%20you%20from%20Doxcela" target="_blank" style="margin: 0 5px;">
                            <img src="https://cdn-icons-png.flaticon.com/512/733/733585.png" alt="WhatsApp" style="height: 24px;">
                        </a>
                    </div>
                    <p style="color: #aaa; font-size: 12px; margin-top: 10px;">&copy; {datetime.now().year} Doxcela. All rights reserved.</p>
                </div>
            </div>
            """

            email_msg = EmailMultiAlternatives(reply_subject, reply_body, 'globalmandateministryinc@gmail.com', [email])
            email_msg.attach_alternative(reply_html, "text/html")
            email_msg.send()

            return JsonResponse({'status': 'success', 'show_modal': True, 'email': email})

        return JsonResponse({'status': 'error', 'message': 'Email required.'})

    # Handle image uploads (non-AJAX form)
    if request.method == 'POST':
        if request.FILES.get('image1'):
            about.image1_url = upload_to_supabase(request.FILES['image1'], request.FILES['image1'].name)
        if request.FILES.get('image2'):
            about.image2_url = upload_to_supabase(request.FILES['image2'], request.FILES['image2'].name)
        if request.FILES.get('image3'):
            about.image3_url = upload_to_supabase(request.FILES['image3'], request.FILES['image3'].name)
        about.save()

    updates = NewUpdate.objects.all().order_by('-upload_date')
    return render(request, 'files/about.html', {
        'about': about,
        'updates': updates
    })
















from django.shortcuts import render
from django.http import JsonResponse
from django.core.mail import EmailMultiAlternatives
from django.views.decorators.csrf import csrf_exempt
from .models import ComingSoonPage, NewsletterSubscriber, UserProfile
from datetime import datetime


def coming_soon(request):
    if request.method == 'POST' and request.headers.get('x-requested-with') == 'XMLHttpRequest':
        email = request.POST.get('email')
        name = 'Subscriber'

        # 🔍 Try to get full name from profile if logged in
        if request.user.is_authenticated:
            profile = UserProfile.objects.filter(user=request.user).first()
            if profile and profile.fullname:
                name = profile.fullname
            elif request.user.get_full_name():
                name = request.user.get_full_name()
            elif request.user.username:
                name = request.user.username

        if email:
            if NewsletterSubscriber.objects.filter(email=email).exists():
                return JsonResponse({
                    'status': 'error',
                    'message': 'Email already subscribed.',
                    'email': email
                })

            NewsletterSubscriber.objects.create(email=email)

            subject = "Thanks for subscribing to our newsletter!"
            text_body = (
                f"Hi {name},\n\n"
                "Thanks for subscribing to our newsletter! "
                "You’re now part of our community and will be the first to receive the latest updates, exclusive content, and special offers.\n\n"
                "We’re excited to have you with us!"
            )
            html_body = text_body.replace('\n', '<br>')

            html_email = f"""
            <div style="font-family: Arial, sans-serif; max-width: 600px; margin: auto; background: #ffffff; border-radius: 10px; padding: 30px; border: 1px solid #eee;">
                <div style="text-align: center; margin-bottom: 20px;">
                    <img src="https://i.imgur.com/yEFgd2V.png" alt="Admin Logo" style="height: 60px;">
                </div>
                <h2 style="color: #1d3557;">Hello {name},</h2>
                <p style="font-size: 16px; color: #333;">{html_body}</p>
                <div style="text-align: center; margin: 30px 0;">
                    <a href="https://gmmi.up.railway.app" style="background: #1d3557; color: white; text-decoration: none; padding: 12px 25px; border-radius: 30px; font-weight: bold;">
                        Visit Our Website
                    </a>
                </div>
                <hr style="border: none; border-top: 1px solid #eee;">
                <div style="text-align: center; margin-top: 20px;">
                    <p style="color: #888; font-size: 14px;">Stay connected with us</p>
                    <div>
                        <a href="https://facebook.com/yourpage" style="margin: 0 5px;">
                            <img src="https://cdn-icons-png.flaticon.com/512/733/733547.png" alt="Facebook" style="height: 24px;">
                        </a>
                        <a href="https://twitter.com/yourpage" style="margin: 0 5px;">
                            <img src="https://cdn-icons-png.flaticon.com/512/733/733579.png" alt="Twitter" style="height: 24px;">
                        </a>
                        <a href="https://instagram.com/yourpage" style="margin: 0 5px;">
                            <img src="https://cdn-icons-png.flaticon.com/512/2111/2111463.png" alt="Instagram" style="height: 24px;">
                        </a>
                        <a href="https://wa.me/2349057147497?text=Hi%2C%20I%20am%20contacting%20you%20from%20Doxcela" target="_blank" style="margin: 0 5px;">
                            <img src="https://cdn-icons-png.flaticon.com/512/733/733585.png" alt="WhatsApp" style="height: 24px;">
                        </a>
                    </div>
                    <p style="color: #aaa; font-size: 12px; margin-top: 10px;">&copy; {datetime.now().year} Doxcela. All rights reserved.</p>
                </div>
            </div>
            """

            email_msg = EmailMultiAlternatives(subject, text_body, 'globalmandateministryinc@gmail.com', [email])
            email_msg.attach_alternative(html_email, "text/html")
            email_msg.send()

            return JsonResponse({'status': 'success', 'show_modal': True, 'email': email})

        return JsonResponse({'status': 'error', 'message': 'Email required.'})

    # GET method
    page = ComingSoonPage.objects.last()

    # Add optional user_has_profile context
    user_has_profile = False
    if request.user.is_authenticated:
        user_has_profile = UserProfile.objects.filter(user=request.user).exists()

    return render(request, 'files/coming_soon.html', {
        'page': page,
        'user_has_profile': user_has_profile
    })


@csrf_exempt
def never_show_modal(request):
    if request.method == 'POST' and request.headers.get('x-requested-with') == 'XMLHttpRequest':
        email = request.POST.get('email')
        if email:
            NewsletterSubscriber.objects.filter(email=email).update(has_closed_modal=True)
            return JsonResponse({'status': 'success'})
    return JsonResponse({'status': 'error'})













from django.shortcuts import render
from django.http import JsonResponse
from django.db.models import Q
from django.templatetags.static import static

from .models import (
    FileUpload, PDFUpload, Hymn, HausaHymn, IgboHymn,
    YorubaHymn, FrenchHymn, ChineseHymn, GermanHymn,
    NewUpdate, UserProfile
)
from .forms import UserProfileForm

# --- SEARCH BOYS ---
def search_api(request):
    query = request.GET.get('q', '').strip()
    results = []

    if request.method == 'POST' and request.headers.get('x-requested-with') == 'XMLHttpRequest':
        if request.user.is_authenticated:
            if UserProfile.objects.filter(user=request.user).exists():
                return JsonResponse({'status': 'error', 'message': "Profile already exists."}, status=400)

            form = UserProfileForm(request.POST, request.FILES)
            if form.is_valid():
                profile = form.save(commit=False)
                profile.user = request.user
                profile.save()
                return JsonResponse({'status': 'success'})
            else:
                return JsonResponse({'status': 'error', 'errors': form.errors}, status=400)
        else:
            return JsonResponse({'status': 'error', 'message': "You must be logged in."}, status=400)

    if query:
        pdf_results = PDFUpload.objects.filter(
            Q(company_location__icontains=query) |
            Q(info_id__icontains=query) |
            Q(date__icontains=query)
        )

        file_results = FileUpload.objects.filter(
            Q(company_location__icontains=query) |
            Q(date__icontains=query) |
            Q(youtube_url__icontains=query)
        )

        hymn_models = [Hymn, HausaHymn, IgboHymn, YorubaHymn, FrenchHymn, ChineseHymn, GermanHymn]
        for model in hymn_models:
            hymn_results = model.objects.filter(
                Q(title__icontains=query) |
                Q(lyrics__icontains=query) |
                Q(description__icontains=query)
            )
            for hymn in hymn_results:
                results.append({
                    "id": hymn.id,
                    "title": hymn.title or "No Title",
                    "description": hymn.description[:100] + "..." if hymn.description else "",
                    "language": model.__name__.replace('Hymn', '') or 'English',
                    "url": f"/search-results/?q={query}&id={hymn.id}&language={model.__name__.replace('Hymn','').lower()}#hymn-{hymn.id}"
                })

        for pdf in pdf_results:
            results.append({
                "id": pdf.id,
                "name": pdf.pdf_url.split('/')[-1] if pdf.pdf_url else "No Name",
                "company_location": pdf.company_location,
                "pdf_url": pdf.pdf_url,
                "image_url": pdf.image_url or 'https://cdn-icons-png.flaticon.com/512/337/337946.png',
                "url": f"/search-results/?q={query}#file-{pdf.id}"
            })

        for file in file_results:
            title = file.company_location or "File"
            results.append({
                "id": file.id,
                "name": title,
                "company_location": file.company_location,
                "youtube_url": file.youtube_url,
                "url": f"/search-results/?q={query}#file-{file.id}"
            })

    user_has_profile = request.user.is_authenticated and UserProfile.objects.filter(user=request.user).exists()
    updates = list(NewUpdate.objects.all().order_by('-upload_date').values('id', 'title', 'upload_date'))

    return JsonResponse({
        "results": results,
        "user_has_profile": user_has_profile,
        "updates": updates
    })
def search_results(request):
    query = request.GET.get('q', '').strip()
    highlight_id = request.GET.get('id')
    highlight_language = request.GET.get('language')

    pdf_results = PDFUpload.objects.filter(
        Q(company_location__icontains=query) |
        Q(info_id__icontains=query) |
        Q(date__icontains=query)
    ) if query else PDFUpload.objects.none()

    file_results = FileUpload.objects.filter(
        Q(company_location__icontains=query) |
        Q(date__icontains=query) |
        Q(youtube_url__icontains=query)
    ) if query else FileUpload.objects.none()

    hymn_models = {
        'english': Hymn,
        'hausa': HausaHymn,
        'igbo': IgboHymn,
        'yoruba': YorubaHymn,
        'french': FrenchHymn,
        'chinese': ChineseHymn,
        'german': GermanHymn,
    }

    hymn_results = Hymn.objects.filter(Q(title__icontains=query) | Q(lyrics__icontains=query) | Q(description__icontains=query)) if query else Hymn.objects.none()
    yoruba_hymn_results = YorubaHymn.objects.filter(Q(title__icontains=query) | Q(lyrics__icontains=query) | Q(description__icontains=query)) if query else YorubaHymn.objects.none()
    hausa_hymn_results = HausaHymn.objects.filter(Q(title__icontains=query) | Q(lyrics__icontains=query) | Q(description__icontains=query)) if query else HausaHymn.objects.none()
    igbo_hymn_results = IgboHymn.objects.filter(Q(title__icontains=query) | Q(lyrics__icontains=query) | Q(description__icontains=query)) if query else IgboHymn.objects.none()
    french_hymn_results = FrenchHymn.objects.filter(Q(title__icontains=query) | Q(lyrics__icontains=query) | Q(description__icontains=query)) if query else FrenchHymn.objects.none()
    chinese_hymn_results = ChineseHymn.objects.filter(Q(title__icontains=query) | Q(lyrics__icontains=query) | Q(description__icontains=query)) if query else ChineseHymn.objects.none()
    german_hymn_results = GermanHymn.objects.filter(Q(title__icontains=query) | Q(lyrics__icontains=query) | Q(description__icontains=query)) if query else GermanHymn.objects.none()

    highlighted_hymns = None
    if highlight_id and highlight_language in hymn_models:
        highlighted_hymns = hymn_models[highlight_language].objects.filter(id=highlight_id)

    user_has_profile = request.user.is_authenticated and UserProfile.objects.filter(user=request.user).exists()

    pdf_results_with_images = []
    for pdf in pdf_results:
        image_url = pdf.image_url if pdf.image_url else static('login-form/images/PDF_image.jpeg')
        pdf_results_with_images.append({
            "id": pdf.id,
            "name": pdf.pdf_url.split('/')[-1] if pdf.pdf_url else "No Name",
            "company_location": pdf.company_location,
            "url": pdf.pdf_url,
            "image": image_url,
            "pdf_file": pdf.pdf_url,
            "date": pdf.date,
            "time": pdf.time,
            "info_id": pdf.info_id
        })

    recent_files = FileUpload.objects.order_by('-date')[:6]
    recent_pdfs = PDFUpload.objects.order_by('-date', '-time')[:6]
    updates = NewUpdate.objects.all().order_by('-upload_date')

    return render(request, "files/search_results.html", {
        "query": query,
        "pdf_results": pdf_results_with_images,
        "file_results": file_results,
        "hymn_results": hymn_results,
        "yoruba_hymn_results": yoruba_hymn_results,
        "hausa_hymn_results": hausa_hymn_results,
        "igbo_hymn_results": igbo_hymn_results,
        "french_hymn_results": french_hymn_results,
        "chinese_hymn_results": chinese_hymn_results,
        "german_hymn_results": german_hymn_results,
        "highlighted_hymns": highlighted_hymns,
        "highlight_language": highlight_language,
        "recent_files": recent_files,
        "recent_pdfs": recent_pdfs,
        "user_has_profile": user_has_profile,
        "updates": updates
    })















@login_required
def send_newsletter(request):
    if request.method == 'POST':
        subject = request.POST.get('subject')
        body = request.POST.get('body')

        # Find all image URLs in the body from the media folder
        image_matches = re.findall(r'src="(/media/[^"]+)"', body)
        attached_images = {}

        for image_url in image_matches:
            image_filename = os.path.basename(image_url)
            image_cid = image_filename  # This will be the Content-ID
            image_path_relative = image_url.replace('/media/', '')  # Remove /media/ from URL
            full_image_path = os.path.join(settings.MEDIA_ROOT, image_path_relative)

            # Replace <img src="..."> with CID and add width & height styling
            pattern = f'src="{image_url}"'
            replacement = f'src="cid:{image_cid}" style="width:600px; height:300px;"'
            body = body.replace(pattern, replacement)

            # Only attach the image if the file actually exists
            if os.path.exists(full_image_path):
                attached_images[image_cid] = full_image_path

        # Send to all newsletter subscribers
        subscribers = NewsletterSubscriber.objects.all()

        for subscriber in subscribers:
            email = EmailMultiAlternatives(
                subject=subject,
                body='This is the fallback plain text version.',
                from_email=settings.DEFAULT_FROM_EMAIL,
                to=[subscriber.email],
            )
            email.attach_alternative(body, "text/html")

            # Attach images as inline MIMEImage
            for cid, path in attached_images.items():
                with open(path, 'rb') as img_file:
                    mime_img = MIMEImage(img_file.read())
                    mime_img.add_header('Content-ID', f'<{cid}>')
                    mime_img.add_header('Content-Disposition', 'inline', filename=cid)
                    email.attach(mime_img)

            email.send(fail_silently=False)

        # Save newsletter record
        DailyNewsletter.objects.create(subject=subject, body=body)

        # Redirect to the same page with a success flag
        return redirect('/send-newsletter/?sent=true')

    # GET request: Render the page with all subscribers
    subscribers = NewsletterSubscriber.objects.all()
    return render(request, 'files/send_newsletter.html', {'subscribers': subscribers})













import uuid
import requests
from django.shortcuts import render, redirect
from django.views.decorators.csrf import csrf_exempt
from django.conf import settings
from .models import Donation

def donation_form(request):
    return render(request, "files/donation_form.html")

def initiate_donation(request):
    if request.method == 'POST':
        name = request.POST.get("name")
        email = request.POST.get("email")
        amount = request.POST.get("amount")
        phone = request.POST.get("phone", "")

        try:
            amount_float = float(amount)
            if amount_float < 100:
                return render(request, "files/error.html", {"message": "Minimum donation is ₦100"})
        except:
            return render(request, "files/error.html", {"message": "Invalid amount"})

        reference = f"PAYSTACK-{uuid.uuid4().hex[:10].upper()}"

        Donation.objects.create(
            name=name,
            email=email,
            phone=phone,
            amount=amount_float,
            reference=reference,
            status="pending"
        )

        headers = {
            "Authorization": f"Bearer {settings.PAYSTACK_SECRET_KEY}",
            "Content-Type": "application/json"
        }

        callback_url = f"{settings.PAYSTACK_CALLBACK_URL}/paystack/verify/"
        return_url = f"{settings.PAYSTACK_CALLBACK_URL}/thank-you/?ref={reference}"

        payload = {
            "email": email,
            "amount": int(amount_float * 100),
            "reference": reference,
            "callback_url": callback_url,
            "return_url": return_url,
        }

        response = requests.post("https://api.paystack.co/transaction/initialize", json=payload, headers=headers)
        data = response.json()

        if data.get("status"):
            auth_url = data["data"]["authorization_url"]
            return redirect(auth_url)
        else:
            return render(request, "files/error.html", {"message": "Paystack Error: Unable to initiate payment"})

    return redirect("donation_form")

@csrf_exempt
def verify_payment(request):
    reference = request.GET.get("reference")
    if not reference:
        return render(request, "files/error.html", {"message": "No reference provided for verification."})

    headers = {
        "Authorization": f"Bearer {settings.PAYSTACK_SECRET_KEY}",
    }

    url = f"https://api.paystack.co/transaction/verify/{reference}"
    response = requests.get(url, headers=headers)
    data = response.json()

    if data.get("status") and data["data"]["status"] == "success":
        try:
            donation = Donation.objects.get(reference=reference)
            donation.status = "paid"
            donation.save()
        except Donation.DoesNotExist:
            pass
        return redirect(f"/thank-you/?ref={reference}")
    else:
        return render(request, "files/error.html", {"message": "Payment verification failed."})

def thank_you(request):
    ref = request.GET.get("ref")
    donation = None
    if ref:
        try:
            donation = Donation.objects.get(reference=ref)
        except Donation.DoesNotExist:
            donation = None

    return render(request, "files/thank_you.html", {"donation": donation})




from django.shortcuts import render
from django.http import JsonResponse
from django.core.mail import send_mail
from django.conf import settings
from .forms import ChildrenMinistryRegistrationForm

def children_ministry(request):
    if request.method == 'POST' and request.headers.get('x-requested-with') == 'XMLHttpRequest':
        form = ChildrenMinistryRegistrationForm(request.POST)
        if form.is_valid():
            instance = form.save()

            subject = "New Children Ministry Registration"
            message = f"""
🧒 CHILDREN MINISTRY REGISTRATION

👤 Parent: {instance.parent_name}
📞 Phone: {instance.phone}
👶 Child: {instance.child_name}
📘 Age Group: {instance.age_group}
📝 Message: {instance.message or 'None'}

🕒 Submitted: {instance.submitted_at.strftime('%Y-%m-%d %I:%M %p')}
"""
            send_mail(subject, message, settings.DEFAULT_FROM_EMAIL, ["globalmandateministryinc@gmail.com"])

            return JsonResponse({'success': True})
        return JsonResponse({'success': False, 'errors': form.errors})

    form = ChildrenMinistryRegistrationForm()
    return render(request, 'files/children_ministry.html', {
        'form': form
    })





from django.shortcuts import render
from .models import WomenEvent

def women_ministry(request):
    events = WomenEvent.objects.all()
    return render(request, 'files/women_ministry.html', {
        'events': events,
    })




from django.utils import timezone
from django.shortcuts import render
from .models import OutreachEvent

def evangelism_ministry(request):
    events = OutreachEvent.objects.filter(date__gte=timezone.now().date())
    return render(request, 'files/evangelism_ministry.html', {
        'events': events,
    })




from .models import YouthEvent

def youth_ministry(request):
    events = YouthEvent.objects.all()
    return render(request, 'files/youth_ministry.html', {
        'events': events,
    })





from django.shortcuts import render
from .models import MenEvent

def men_ministry(request):
    events = MenEvent.objects.all()
    return render(request, 'files/men_ministry.html', {
        'events': events,
    })






from .models import GMSOMSlide

def overview(request):
    slides = GMSOMSlide.objects.all()
    return render(request, 'files/overview.html', {'slides': slides})







from django.http import JsonResponse
from django.core.mail import send_mail
from django.conf import settings
from .forms import PrayerRequestForm

def prayer_request_view(request):
    if request.method == 'POST':
        form = PrayerRequestForm(request.POST)
        if form.is_valid():
            prayer = form.save()

            # ✅ Send confirmation email
            send_mail(
                subject="🙏 Your Prayer Request Was Received",
                message=(
                    f"Dear {prayer.name},\n\n"
                    f"Thank you for your prayer request:\n\n"
                    f"\"{prayer.message}\"\n\n"
                    f"We are standing with you in faith!\n\n"
                    f"📞 Phone: {prayer.phone or 'Not provided'}\n\n"
                    f"Blessings,\nGMMIConnect Team"
                ),
                from_email=settings.DEFAULT_FROM_EMAIL,
                recipient_list=[prayer.email],
                fail_silently=False,
            )
            return JsonResponse({'status': 'success'})
        else:
            return JsonResponse({'status': 'fail', 'errors': form.errors}, status=400)
    return JsonResponse({'status': 'invalid'}, status=405)






# def testimony_view(request):
#     testimonies = Testimony.objects.all().order_by('-created_at')
#     return render(request, 'files/testimony.html', {'testimonies': testimonies})







from django.shortcuts import render
from django.http import JsonResponse, HttpResponse
from django.core.mail import EmailMessage, EmailMultiAlternatives
from django.template.loader import render_to_string
from django.conf import settings
from io import BytesIO
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Inches
from .forms import SchoolOfMinistryForm
import mimetypes


def generate_pdf(context):
    html = render_to_string("files/pdf_template.html", context)
    result = BytesIO()
    pisa.pisaDocument(BytesIO(html.encode("UTF-8")), result)
    return result.getvalue()


def generate_word(context):
    data = context['registration']
    doc = Document()
    doc.add_heading('GMMI School of Ministry Registration', 0)

    doc.add_paragraph(f"Full Name: {data.full_name}")
    doc.add_paragraph(f"Email: {data.email}")
    doc.add_paragraph(f"Phone: {data.phone}")
    doc.add_paragraph(f"Gender: {data.gender}")
    doc.add_paragraph(f"Date of Birth: {data.dob}")
    doc.add_paragraph(f"Marital Status: {data.marital_status}")
    doc.add_paragraph(f"Country: {data.country}")
    doc.add_paragraph(f"State/City: {data.state_city}")
    doc.add_paragraph(f"Occupation: {data.occupation}")
    doc.add_paragraph(f"Church: {data.church}")
    doc.add_paragraph(f"Born Again: {data.born_again}")
    doc.add_paragraph(f"Baptized in Holy Ghost: {data.holy_ghost}")
    doc.add_paragraph(f"Reason: {data.reason}")

    # Add student photo if available
    if data.photo:
        try:
            doc.add_paragraph("Passport Photograph:")
            doc.add_picture(data.photo.path, width=Inches(1.5))
        except Exception as e:
            doc.add_paragraph("Photo could not be loaded.")

    word_stream = BytesIO()
    doc.save(word_stream)
    word_stream.seek(0)
    return word_stream


def register(request):
    if request.method == 'POST':
        form = SchoolOfMinistryForm(request.POST, request.FILES)
        if form.is_valid():
            instance = form.save()
            context = {'registration': instance}

            # Generate PDF and Word
            pdf_file = generate_pdf(context)
            word_file = generate_word(context)

            # Prepare Admin Email
            admin_email = EmailMessage(
                subject="New GMMI Registration",
                body=f"A new registration has been submitted by {instance.full_name}.",
                from_email=settings.DEFAULT_FROM_EMAIL,
                to=["globalmandateministryinc@gmail.com"],
            )

            # Attach PDF
            if pdf_file:
                admin_email.attach(
                    f"{instance.full_name}_registration.pdf",
                    pdf_file,
                    "application/pdf"
                )

            # Attach Word Doc
            if word_file:
                admin_email.attach(
                    f"{instance.full_name}_registration.docx",
                    word_file.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            # Attach Image Safely
            if instance.photo:
                mime_type, _ = mimetypes.guess_type(instance.photo.name)
                admin_email.attach(
                    instance.photo.name,
                    instance.photo.read(),
                    mime_type or 'image/jpeg'
                )

            admin_email.send(fail_silently=False)

            # ✉️ Send Confirmation Email to Student (with embedded logo)
            html_message = render_to_string("files/email_to_student.html", {"name": instance.full_name})

            student_email = EmailMultiAlternatives(
                subject="🎓 GMMI Registration Successful!",
                body="Thank you for registering. Please view this email in HTML.",
                from_email=settings.DEFAULT_FROM_EMAIL,
                to=[instance.email],
            )
            student_email.attach_alternative(html_message, "text/html")

            # Embed GMMI logo image inside the email
            logo_path = os.path.join(settings.BASE_DIR, "static", "login-form", "images", "GMMI_LOGO.png")
            if os.path.exists(logo_path):
                with open(logo_path, 'rb') as f:
                    logo = MIMEImage(f.read())
                    logo.add_header('Content-ID', '<gmmi_logo>')
                    logo.add_header('Content-Disposition', 'inline', filename="GMMI_LOGO.png")
                    student_email.attach(logo)

            student_email.send(fail_silently=False)

            return JsonResponse({'success': True})
        else:
            return JsonResponse({'success': False, 'errors': form.errors})

    form = SchoolOfMinistryForm()
    return render(request, 'files/register.html', {'form': form})
